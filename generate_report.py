"""
Generate the AlphaEire Capital CEAI CA Assignment Word report.
Run: python generate_report.py
Output: AlphaEire_Capital_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def add_para(doc, text, bold=False, italic=False, size=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_label_text(doc, label, text):
    """Bold label followed by normal text in same paragraph."""
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(3)
    return p


def section_divider(doc):
    doc.add_paragraph()


GREEN = (0, 200, 150)
DARK_GREY = (80, 80, 80)

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ───────────────────────────────────────────────────────────────

title = doc.add_heading('AlphaEire Capital', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 200, 150)

sub = doc.add_heading('Agentic Investment Pipeline', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('CEAI CA Assignment\n').bold = True
meta.add_run(f'Submitted: {datetime.date.today().strftime("%B %Y")}\n')
meta.add_run('GitHub Pages: https://vinayakcsnci.github.io/alphaeire-capital/')

doc.add_page_break()

# ── SECTION 1: YOUR ORGANISATION ─────────────────────────────────────────────

set_heading(doc, '1. Your Organisation', 1, GREEN)

add_para(doc, (
    'AlphaEire Capital is a fictional Irish investment firm that specialises exclusively in '
    'ISEQ (Irish Stock Exchange Quotient)-listed equities. The firm\'s core challenge is that '
    'institutional-grade market intelligence — systematic stock screening, real-time alerting, '
    'professional investor communications, and strategic review — typically requires a full team '
    'of analysts, designers, developers, and communications professionals. For smaller boutique '
    'fund managers and retail investors operating in the Irish market, this expertise is simply '
    'out of reach.'
))

add_para(doc, (
    'The agentic solution collapses an entire investment intelligence workflow — from raw market '
    'analysis through dashboard design, working prototype, investor communications, and executive '
    'strategy — into a five-agent automated pipeline that completes in minutes. Where a human '
    'team would take days to scan stocks, commission dashboard designs, build prototypes, draft '
    'investor emails, and produce management reports, AlphaEire\'s pipeline does it sequentially '
    'and automatically.'
))

add_para(doc, (
    'The agentic structure is a natural fit because each phase of this workflow requires '
    'genuinely different expertise: equity analysis, UX design, software development, marketing '
    'copywriting, and executive strategy cannot be collapsed into a single generalised agent '
    'without losing quality. The five-agent pipeline mirrors how a real investment firm organises '
    'its people — each specialist acting on the prior team member\'s output, amplifying it into '
    'something no single agent could have produced alone. The handoff is the heart of the system.'
))

section_divider(doc)

# ── SECTION 2: AGENT DESIGNS ─────────────────────────────────────────────────

set_heading(doc, '2. Agent Designs', 1, GREEN)

add_para(doc, (
    'Each agent has a distinct name, role, full system prompt, defined output, and a deliberately '
    'differentiated personality. The five archetypes — Researcher, Designer, Maker, Communicator, '
    'and Manager — are not generic chatbots; they are domain experts whose system prompts shape '
    'their vocabulary, output structure, and professional register.'
), space_after=10)

# Agent 1
set_heading(doc, 'Agent 1 — Aoife (Researcher)', 2)
add_label_text(doc, 'Role', 'Senior Equity Analyst')
add_label_text(doc, 'Produces', 'Structured research brief with 3–5 ISEQ stock picks grounded in live Yahoo Finance data: signal types, live-price rationale, risk flags, and market outlook')
add_label_text(doc, 'Live Data', '^ISEQ, KRX.IR, KRZ.IR, GL9.IR, IR5B.IR, EG7.IR, IRES.IR, OIZ.IR, HSW.IR, MIO.IR, MLC.IR, 8GW.IR — fetched from Yahoo Finance (Euronext Dublin) before each pipeline run')

add_para(doc, 'System Prompt:', bold=True)
add_para(doc, (
    '"You are Aoife, a senior equity analyst at AlphaEire Capital, an Irish investment firm focused '
    'on ISEQ-listed stocks. Live market data from Yahoo Finance will be provided at the top of the '
    'user message — treat it as ground truth for current prices and movements. Your role is to analyse '
    'this live ISEQ data, identify 3-5 stocks showing strong signals (momentum, undervaluation, or '
    'volatility opportunity), and produce a structured research brief. For each stock provide: ticker, '
    'sector, current price, signal type, key rationale (2-3 sentences grounded in the live data), and '
    'a risk flag. Close with an overall Irish market outlook paragraph. Be precise, data-driven, and professional."'
), italic=True)
add_para(doc, (
    'Data pipeline: Before Aoife runs, the system fetches real-time prices, day ranges, and 52-week '
    'ranges for 7 ISEQ symbols from the Yahoo Finance v8/chart API (browser) or the yfinance Python '
    'library (pipeline.py). This data is prepended to her user message, grounding every stock pick '
    'in current market reality rather than training-data estimates.'
))
add_para(doc, (
    'Personality: Analytically rigorous and evidence-first. Aoife speaks with the measured authority '
    'of a senior CFA charterholder — structured output, no speculation, every claim referenced to '
    'a live data point. Her output is the foundation the entire pipeline builds on.'
))

doc.add_paragraph()

# Agent 2
set_heading(doc, 'Agent 2 — Ciarán (Designer)', 2)
add_label_text(doc, 'Role', 'Product Designer')
add_label_text(doc, 'Produces', 'Dashboard design specification covering layout, alert trigger rules, data visualisation choices, and UX rationale')

add_para(doc, 'System Prompt:', bold=True)
add_para(doc, (
    '"You are Ciarán, a product designer at AlphaEire Capital. You have received a research brief '
    'from our equity analyst team. Design a stock monitoring dashboard and real-time alert system '
    'for our portfolio managers. Produce a detailed design specification covering: (1) dashboard '
    'layout and key UI components with exact descriptions, (2) alert trigger rules for each stock '
    'signal type identified in the research, (3) data visualisation choices (chart types, colour '
    'coding, thresholds), (4) UX rationale for every major decision. Format your output as a '
    'structured design specification document with clear headings."'
), italic=True)
add_para(doc, (
    'Personality: Creative and systematic. Ciarán frames every decision around the portfolio '
    'manager\'s cognitive workflow, not aesthetics. His specifications are concrete enough that '
    'a developer can implement them without follow-up questions.'
))

doc.add_paragraph()

# Agent 3
set_heading(doc, 'Agent 3 — Siobhán (Maker)', 2)
add_label_text(doc, 'Role', 'Senior Full-Stack Developer')
add_label_text(doc, 'Produces', 'Working self-contained HTML/CSS/JS stock monitoring dashboard with Chart.js charts, signal badges, and alert panel')

add_para(doc, 'System Prompt:', bold=True)
add_para(doc, (
    '"You are Siobhán, a senior full-stack developer at AlphaEire Capital. You have received a '
    'design specification for a stock monitoring dashboard. Build a working, self-contained HTML '
    'file (with embedded CSS and JavaScript) that implements this design. Use realistic placeholder '
    'data for the ISEQ stocks identified in the research brief. The dashboard must include: stock '
    'cards with signal badges (BUY/WATCH/ALERT), a line chart using Chart.js (load via CDN: '
    'https://cdn.jsdelivr.net/npm/chart.js), and an alert notification panel on the right side. '
    'Output ONLY the complete HTML file contents — no explanations before or after, just the raw '
    'HTML starting with <!DOCTYPE html>."'
), italic=True)
add_para(doc, (
    'Personality: Pragmatic and precise. Siobhán prioritises functional correctness over perfection, '
    'delivering a working prototype on the shortest path. Her output is the tangible artefact that '
    'makes the entire pipeline visible — Siobhán\'s dashboard is the landing page of the site.'
))

doc.add_paragraph()

# Agent 4
set_heading(doc, 'Agent 4 — Declan (Communicator)', 2)
add_label_text(doc, 'Role', 'Head of Investor Relations')
add_label_text(doc, 'Produces', 'Professional investor alert email and two LinkedIn posts (retail + institutional)')

add_para(doc, 'System Prompt:', bold=True)
add_para(doc, (
    '"You are Declan, head of investor relations at AlphaEire Capital. Based on the equity research '
    'brief and the AI-powered monitoring dashboard built by our technology team, produce two things: '
    '(1) A professional investor alert email with subject line and body (~250 words) announcing the '
    'top Irish stock opportunities to our client base. (2) Two LinkedIn posts — one targeting Irish '
    'retail investors (casual, engaging, ~100 words) and one targeting institutional investors (formal, '
    'data-led, ~120 words) — both promoting AlphaEire Capital\'s new AI-powered ISEQ monitoring '
    'capability. Tone: confident, trustworthy, Irish-market-focused."'
), italic=True)
add_para(doc, (
    'Personality: Confident, persuasive, and trust-aware. Declan never over-promises returns. '
    'His communications build credibility through specificity — readers should feel informed and '
    'confident, not sold to.'
))

doc.add_paragraph()

# Agent 5
set_heading(doc, 'Agent 5 — Margaret (Manager)', 2)
add_label_text(doc, 'Role', 'Chief Investment Officer')
add_label_text(doc, 'Produces', 'Executive summary, 90-day operational plan (5 milestones), and Risk & Compliance section')

add_para(doc, 'System Prompt:', bold=True)
add_para(doc, (
    '"You are Margaret, Chief Investment Officer at AlphaEire Capital. Review the full pipeline '
    'output from your team: the equity research brief, the dashboard design specification, the '
    'working prototype description, and the investor communications. Produce three sections: '
    '(1) Executive Summary (~150 words) assessing strategic alignment and output quality. '
    '(2) 90-Day Operational Plan with exactly 5 concrete, numbered milestones. '
    '(3) Risk & Compliance section covering: GDPR implications of using AI for investment '
    'recommendations, EU AI Act classification of this system, and how AlphaEire Capital will '
    'maintain customer trust. Be authoritative, strategic, and candid about both strengths and gaps."'
), italic=True)
add_para(doc, (
    'Personality: Authoritative and candid. Margaret is the only agent who can critique the others\' '
    'work — she sees the full picture. She writes in the register of a seasoned CIO: no hedging, '
    'no filler, no praise for its own sake.'
))

section_divider(doc)

# ── SECTION 3: PIPELINE IN ACTION ────────────────────────────────────────────

set_heading(doc, '3. The Pipeline in Action', 1, GREEN)

add_para(doc, (
    'Before the pipeline starts, live market data is fetched from Yahoo Finance for 7 ISEQ symbols '
    '(^ISEQ All Share index, AIBG.I, BIRG.I, RYA.I, CRH.L, DCC.L, PTSB.I): current price, '
    'day range, and 52-week range. This data is prepended to Aoife\'s user message so her stock '
    'picks are grounded in current reality, not training-data estimates.'
))
add_para(doc, (
    'Each agent\'s complete text output is then injected directly into the next agent\'s user '
    'message as structured context, delimited by labelled section headers '
    '(e.g. --- RESEARCH BRIEF ---). Every downstream agent has access to the full upstream '
    'reasoning, not a summarised version of it.'
))

set_heading(doc, 'Handoff Chain', 2)
steps = [
    ('Aoife → Ciarán', 'Research brief (3–5 stocks, signals, rationale, market outlook) → Design specification'),
    ('Ciarán → Siobhán', 'Design spec (layout, alert rules, chart types, UX rationale) → Working HTML dashboard'),
    ('Siobhán → Declan', 'Confirmation of dashboard build + research brief → Investor email + LinkedIn posts'),
    ('Declan → Margaret', 'All four prior outputs → Executive summary + 90-day plan + compliance review'),
]
for label, desc in steps:
    p = doc.add_paragraph()
    p.add_run(label + ': ').bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

set_heading(doc, 'Demo Modes', 2)
add_para(doc, (
    'The GitHub Pages site runs the pipeline in two modes:'
))
add_bullet(doc, (
    'Demo Mode: Pre-generated outputs are loaded from demo_data.js. No API key required. '
    'All five agent cards show "Done" immediately and each output is viewable by clicking the card.'
))
add_bullet(doc, (
    'Live Mode: Supports three AI providers — Anthropic (Claude claude-opus-4-7), OpenAI (GPT-4o), '
    'and Groq (Llama 3.3 70B). The user enters their API key, selects a provider, and clicks '
    '"Run Full Pipeline". Each agent card animates Running → Done sequentially and the output '
    'appears live in the panel.'
))

add_para(doc, '')
add_label_text(doc, 'GitHub Pages', 'https://vinayakcsnci.github.io/alphaeire-capital/')
add_label_text(doc, 'GitHub Repository', 'https://github.com/Vinayakcsnci/alphaeire-capital')

set_heading(doc, 'Evidence of Output', 2)
add_para(doc, (
    'Siobhán\'s HTML dashboard is embedded directly as the landing page — when you visit the '
    'site, you see the working ISEQ monitoring dashboard immediately, full-screen, built by the '
    'Maker agent. The "Open Agent Pipeline →" button navigates to the pipeline UI where each '
    'agent\'s full text output is viewable. Margaret\'s executive report, Aoife\'s research brief, '
    'Declan\'s investor email — all are accessible and readable in the output panel.'
))

section_divider(doc)

# ── SECTION 4: REGULATORY AND ETHICAL ────────────────────────────────────────

set_heading(doc, '4. Regulatory and Ethical Considerations', 1, GREEN)

set_heading(doc, 'GDPR', 2)
add_para(doc, (
    'In its current form, the AlphaEire Capital pipeline processes no personal data about '
    'end-users — it analyses publicly available market information only. The Anthropic, OpenAI, '
    'and Groq API keys are entered client-side in the browser and are never transmitted to or '
    'stored on any AlphaEire server, consistent with the GDPR\'s data minimisation principle '
    '(Article 5(1)(c)).'
))
add_para(doc, (
    'If the system were extended to personalise recommendations by user portfolio or behavioural '
    'profile, additional obligations would arise under Articles 13–22: lawful basis for processing, '
    'right to explanation for automated decisions with significant effects (Article 22), and data '
    'subject access rights. Investment recommendations generated by an AI system that influences '
    'financial decisions could be classified as "solely automated processing" under Article 22, '
    'triggering the right to human review on request.'
))

set_heading(doc, 'EU AI Act', 2)
add_para(doc, (
    'The EU AI Act (Regulation 2024/1689, in force August 2024) classifies AI systems by risk. '
    'AlphaEire Capital\'s pipeline operates in the financial services domain, where AI used to '
    'influence investment decisions could be classified as High-Risk under Annex III, Article 5b '
    '(AI systems that affect access to financial services or evaluate creditworthiness). At minimum, '
    'it falls under the General-Purpose AI (GPAI) provisions of Title VII, requiring: transparency '
    'to users that outputs are AI-generated; technical documentation describing system capabilities '
    'and limitations; and compliance with the EU copyright acquis for training data.'
))
add_para(doc, (
    'AlphaEire Capital explicitly positions the pipeline as an intelligence tool, not a regulated '
    'investment adviser. All outputs carry the implicit requirement to be reviewed by a licensed '
    'financial professional before any client-facing action is taken. This positioning — as a '
    'decision-support system, not a decision-making system — is the principal mechanism for '
    'managing AI Act risk classification.'
))

set_heading(doc, 'Customer Trust', 2)
add_para(doc, (
    'The five-agent pipeline creates an inherent audit trail — each agent\'s reasoning is visible '
    'and reviewable on the Pipeline page of the site. A user can read exactly what Aoife found, '
    'what Ciarán designed, what Siobhán built, and what Margaret assessed. This transparency is '
    'itself a trust signal: the system cannot hide its reasoning behind an opaque recommendation.'
))
add_para(doc, (
    'All investor communications produced by Declan include standard investment disclaimers. '
    'Margaret\'s compliance section explicitly flags the AI-generated nature of the outputs and '
    'the requirement for professional review. Building these guardrails into the agents\' system '
    'prompts — rather than bolting them on externally — means they are structurally present in '
    'every pipeline run.'
))

section_divider(doc)

# ── SECTION 5: REFLECTION ────────────────────────────────────────────────────

set_heading(doc, '5. Reflection', 1, GREEN)

set_heading(doc, 'What Worked', 2)
add_para(doc, (
    'The pipeline handoff mechanism exceeded expectations. Passing the full text of each agent\'s '
    'output as structured context to the next agent — with explicit section headers — created '
    'genuinely cumulative results. Siobhán\'s dashboard included the exact tickers Aoife '
    'identified; Declan\'s LinkedIn post referenced the specific momentum signals in the research '
    'brief; Margaret\'s 90-day plan named the features Siobhán had actually built. The agents '
    'built on each other\'s work in ways that felt emergent rather than scripted.'
))
add_para(doc, (
    'The multi-provider architecture (Anthropic / OpenAI / Groq) also worked well — the same '
    'pipeline runs identically regardless of which LLM powers it, because the agent logic lives '
    'in the system prompts, not in provider-specific features.'
))

set_heading(doc, 'What Did Not Work', 2)
add_para(doc, (
    'The Maker agent (Siobhán) proved the most fragile. HTML generation at 3,000 tokens can '
    'produce incomplete output if the design specification is long, occasionally cutting off '
    'before the closing </html> tag. Requesting the output start with <!DOCTYPE html> and '
    'pre-processing it in the iframe\'s srcdoc attribute mitigated but did not eliminate this. '
    'A production system would use streaming with a token buffer and a fallback truncation handler.'
))

set_heading(doc, 'What Surprised Me', 2)
add_para(doc, (
    'How much personality the agents developed purely from system prompt design. Aoife\'s output '
    'reads measurably differently from Margaret\'s — different sentence structure, different '
    'vocabulary register, different level of assertiveness — not because of any special mechanism, '
    'but because their system prompts shape tone at the level of word choice. The five agents '
    'feel like distinct professionals, not five instances of the same model running different '
    'task prompts.'
))
add_para(doc, (
    'The quality of the handoff context design mattered far more than expected. An ambiguous '
    'handoff (e.g. "the developer built something useful") produced noticeably weaker downstream '
    'output than a structured one with labelled sections and explicit content summaries. The '
    'pipeline quality scales directly with the precision of the inter-agent context design.'
))

set_heading(doc, 'What I Learned', 2)
add_para(doc, (
    'Multi-agent collaboration is not about running multiple agents in parallel — it is about '
    'designing the information flow between them. Each handoff boundary is an engineering decision: '
    'what does the next agent need, in what format, to do its best work? Getting this wrong at '
    'any single handoff degrades all subsequent outputs. Getting it right compounds positively '
    'through the chain.'
))
add_para(doc, (
    'The five archetypes in the assignment brief (Researcher, Designer, Maker, Communicator, '
    'Manager) map almost perfectly onto a real investment intelligence workflow. This is not '
    'coincidental — the archetypes represent genuinely distinct cognitive modes. An agentic '
    'organisation works best when its agent boundaries follow natural workflow boundaries, '
    'not arbitrary task splits.'
))

set_heading(doc, 'Future Improvements', 2)
add_bullet(doc, (
    'Streaming output: Display each agent\'s response word-by-word in the UI as it is generated, '
    'rather than waiting for the full response before rendering.'
))
add_bullet(doc, (
    'Persistent memory: Allow Margaret to compare this pipeline run\'s findings with previous runs, '
    'tracking stock signal changes over time.'
))
add_bullet(doc, (
    'Error recovery: If the Maker agent produces malformed HTML, automatically retry with a '
    'simplified design specification rather than halting the entire pipeline.'
))
add_bullet(doc, (
    'Real market data: Integrate with a free Irish market data API (e.g. Yahoo Finance) so '
    'Aoife\'s research brief reflects actual current prices and volumes, not simulated data.'
))
add_bullet(doc, (
    'PDF report export: Add a button that packages all five agent outputs into a single branded '
    'PDF report for investor distribution.'
))

section_divider(doc)

# ── FOOTER NOTE ──────────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'AlphaEire Capital · Built with Claude claude-opus-4-7 · ISEQ Market Intelligence · '
    + datetime.date.today().strftime('%B %Y')
)
run.font.color.rgb = RGBColor(139, 148, 158)
run.font.size = Pt(9)

# ── SAVE ─────────────────────────────────────────────────────────────────────

output_path = 'AlphaEire_Capital_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
